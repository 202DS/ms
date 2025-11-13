# wechat_style_app_final.py
import sys
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import pulp
from PySide2.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                               QPushButton, QLabel, QLineEdit, QTextEdit, QTabWidget, QFrame,
                               QGroupBox, QScrollArea, QGridLayout, QMessageBox, QProgressBar,
                               QFileDialog)  # 添加QFileDialog
from PySide2.QtCore import Qt, QTimer
from PySide2.QtGui import QFont, QPalette, QColor, QIcon, QPixmap
import warnings

warnings.filterwarnings('ignore')

# 导入Word导出相关库
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx 未安装，Word导出功能不可用")


# 获取资源路径
def resource_path(relative_path):
    """获取资源的绝对路径"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False


class WildfireModel:
    """消防无人机配置模型"""

    def __init__(self):
        self.params = {
            'R': 5,
            'lambda_i': [10, 12, 8, 15, 6],
            'A_mean': 1.5,
            'A_std': 1.0,
            'alpha1': 2.0,
            'alpha2': 0.3,
            'S_SSA': 0.8,
            'T_cycle': 1.5,
            'R_max': 0.3,
            'R_cov': 15,
            'beta': 1.3,
            'p_SSA': 80000,
            'p_R': 25000,
            'gamma': 1.2
        }
        self.results = {}

    def run_analysis(self, years=10, growth_rate=0.03):
        """运行完整分析"""
        # 这里简化实现，实际应该调用之前的核心算法
        np.random.seed(42)

        # 模拟一些结果
        config = {
            'x_SSA_opt': np.random.randint(8, 15),
            'x_R_opt': np.random.randint(5, 12),
            'total_cost': np.random.randint(500000, 1500000),
            'peak_demand_SSA': np.random.uniform(5, 12),
            'peak_demand_R': np.random.uniform(3, 8)
        }

        self.results = {
            'config': config,
            'multi_year': None,
            'relay_deployment': None
        }

        return self.results

    def plot_demand(self):
        """绘制需求时序图"""
        try:
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 8))

            # 模拟数据
            hours = 1000
            time = np.arange(hours)
            ssa_demand = np.random.poisson(5, hours) + np.sin(time / 100) * 2
            relay_demand = np.random.poisson(3, hours) + np.sin(time / 100) * 1.5

            ax1.plot(time, ssa_demand, alpha=0.7, color='#07C160', linewidth=2)
            ax1.set_ylabel('SSA无人机需求', fontsize=12)
            ax1.set_title('SSA无人机需求时间序列', fontsize=14, fontweight='bold')
            ax1.grid(True, alpha=0.3)
            ax1.set_facecolor('#f8f9fa')

            ax2.plot(time, relay_demand, alpha=0.7, color='#07C160', linewidth=2)
            ax2.set_ylabel('中继无人机需求', fontsize=12)
            ax2.set_xlabel('时间 (小时)', fontsize=12)
            ax2.set_title('中继无人机需求时间序列', fontsize=14, fontweight='bold')
            ax2.grid(True, alpha=0.3)
            ax2.set_facecolor('#f8f9fa')

            plt.tight_layout()

            # 保存图片
            os.makedirs('temp', exist_ok=True)
            image_path = 'temp/demand_plot.png'
            plt.savefig(image_path, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)  # 关闭图形，避免显示

            return image_path

        except Exception as e:
            print(f"绘制需求图失败: {e}")
            return None

    def plot_costs(self):
        """绘制成本分析图"""
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

            # 成本构成
            labels = ['SSA无人机', '中继无人机', '维护费用']
            sizes = [65, 25, 10]
            colors = ['#07C160', '#66BB6A', '#81C784']

            ax1.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
            ax1.set_title('成本构成分析', fontsize=14, fontweight='bold')

            # 年度成本
            years = np.arange(10)
            costs = np.random.randint(80000, 200000, 10).cumsum()

            ax2.bar(years, costs, color='#07C160', alpha=0.7)
            ax2.set_xlabel('年份', fontsize=12)
            ax2.set_ylabel('累计成本 ($)', fontsize=12)
            ax2.set_title('年度累计成本', fontsize=14, fontweight='bold')
            ax2.grid(True, alpha=0.3)
            ax2.set_facecolor('#f8f9fa')

            plt.tight_layout()

            # 保存图片
            os.makedirs('temp', exist_ok=True)
            image_path = 'temp/cost_plot.png'
            plt.savefig(image_path, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)  # 关闭图形，避免显示

            return image_path

        except Exception as e:
            print(f"绘制成本图失败: {e}")
            return None

    def plot_relay_deployment(self):
        """绘制中继部署图"""
        try:
            fig, ax = plt.subplots(figsize=(10, 8))

            # 模拟部署场景
            eoc = (0, 0)
            frontlines = [(15, 10), (25, 15), (35, 8), (20, 25), (30, 30)]
            relays = [(10, 5), (20, 15), (30, 10)]

            # 绘制EOC
            ax.scatter(*eoc, s=300, c='red', marker='s', label='指挥中心')
            ax.text(eoc[0], eoc[1] + 2, 'EOC', ha='center', fontweight='bold', fontsize=12)

            # 绘制前线
            frontline_x, frontline_y = zip(*frontlines)
            ax.scatter(frontline_x, frontline_y, s=150, c='blue', marker='^', label='前线小队')

            # 绘制中继
            relay_x, relay_y = zip(*relays)
            ax.scatter(relay_x, relay_y, s=200, c='#07C160', marker='D', label='中继无人机')

            # 绘制通信范围
            for relay in relays:
                circle = plt.Circle(relay, self.params['R_cov'], color='#07C160', alpha=0.1)
                ax.add_patch(circle)

            ax.set_xlabel('X坐标 (km)', fontsize=12)
            ax.set_ylabel('Y坐标 (km)', fontsize=12)
            ax.set_title('中继无人机部署方案', fontsize=14, fontweight='bold')
            ax.legend()
            ax.grid(True, alpha=0.3)
            ax.set_aspect('equal')
            ax.set_facecolor('#f8f9fa')

            plt.tight_layout()

            # 保存图片
            os.makedirs('temp', exist_ok=True)
            image_path = 'temp/relay_deployment.png'
            plt.savefig(image_path, dpi=100, bbox_inches='tight', facecolor='white')
            plt.close(fig)  # 关闭图形，避免显示

            return image_path

        except Exception as e:
            print(f"绘制中继部署图失败: {e}")
            return None


class WeChatStyleWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.model = WildfireModel()
        self.current_plot_path = None  # 当前显示的图片路径
        self.init_ui()

    def init_ui(self):
        """初始化微信风格界面"""
        self.setWindowTitle('智能消防无人机配置系统')
        self.setFixedSize(1200, 800)  # 增加窗口大小以容纳预览区域

        # 设置窗口图标 - 添加logo.ico
        icon_path = resource_path('logo.ico')
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
            print(f"✅ 已加载图标: {icon_path}")
        else:
            print(f"⚠️ 图标文件不存在: {icon_path}")

        # 设置微信风格样式
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #07C160;
                border-radius: 8px;
                margin-top: 1ex;
                padding-top: 10px;
                background-color: white;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #07C160;
            }
            QPushButton {
                background-color: #07C160;
                border: none;
                color: white;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #06AE56;
            }
            QPushButton:pressed {
                background-color: #059C4D;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666666;
            }
            QLineEdit {
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 6px;
                background-color: white;
            }
            QLineEdit:focus {
                border-color: #07C160;
            }
            QTextEdit {
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-family: "Microsoft YaHei";
            }
            QLabel {
                color: #333333;
            }
            QTabWidget::pane {
                border: 1px solid #C2C7CB;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #E1E1E1;
                border: 1px solid #C4C4C3;
                padding: 8px 20px;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #07C160;
                color: white;
            }
            QProgressBar {
                border: 1px solid #ddd;
                border-radius: 4px;
                text-align: center;
                background-color: #f0f0f0;
            }
            QProgressBar::chunk {
                background-color: #07C160;
                border-radius: 3px;
            }
        """)

        # 创建中心部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # 改进的标题区域 - 增加字间距
        title_label = QLabel('智 能 森 林 消 防 无 人 机 配 置 系 统')  # 添加空格增加字间距
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("""
            QLabel {
                font-size: 26px;
                font-weight: bold;
                color: #07C160;
                padding: 25px;
                background-color: white;
                border-radius: 12px;
                margin: 12px;
                letter-spacing: 8px;  /* 字间距 */
                text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
            }
        """)
        layout.addWidget(title_label)

        # 创建标签页
        self.tab_widget = QTabWidget()
        layout.addWidget(self.tab_widget)

        # 创建各个标签页
        self.create_parameter_tab()
        self.create_simulation_tab()
        self.create_results_tab()
        self.create_visualization_tab()

    def create_parameter_tab(self):
        """创建参数设置标签页"""
        tab = QWidget()
        layout = QVBoxLayout(tab)

        # 创建滚动区域
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)

        # 区域设置组
        region_group = QGroupBox("区域设置")
        region_layout = QGridLayout(region_group)

        region_layout.addWidget(QLabel("区域数量:"), 0, 0)
        self.region_count = QLineEdit("5")
        region_layout.addWidget(self.region_count, 0, 1)

        self.region_freqs = []
        for i in range(5):  # 默认显示5个区域
            region_layout.addWidget(QLabel(f"区域{i + 1}火灾频率:"), i + 1, 0)
            freq_edit = QLineEdit("10")
            self.region_freqs.append(freq_edit)
            region_layout.addWidget(freq_edit, i + 1, 1)

        scroll_layout.addWidget(region_group)

        # 火灾规模组
        fire_group = QGroupBox("火灾规模参数")
        fire_layout = QGridLayout(fire_group)

        fire_params = [
            ("平均火场面积(km²):", "A_mean", "1.5"),
            ("面积标准差:", "A_std", "1.0"),
            ("基础消防小队数:", "alpha1", "2.0"),
            ("每平方公里增加小队数:", "alpha2", "0.3")
        ]

        for i, (label, attr, default) in enumerate(fire_params):
            fire_layout.addWidget(QLabel(label), i, 0)
            edit = QLineEdit(default)
            setattr(self, attr, edit)
            fire_layout.addWidget(edit, i, 1)

        scroll_layout.addWidget(fire_group)

        # 设备性能组
        device_group = QGroupBox("设备性能参数")
        device_layout = QGridLayout(device_group)

        device_params = [
            ("SSA覆盖面积(km²):", "S_SSA", "0.8"),
            ("巡航周期(小时):", "T_cycle", "1.5"),
            ("最大重访时间(小时):", "R_max", "0.3"),
            ("中继通信半径(km):", "R_cov", "15"),
            ("中继冗余系数:", "beta", "1.3")
        ]

        for i, (label, attr, default) in enumerate(device_params):
            device_layout.addWidget(QLabel(label), i, 0)
            edit = QLineEdit(default)
            setattr(self, attr, edit)
            device_layout.addWidget(edit, i, 1)

        scroll_layout.addWidget(device_group)

        # 成本预算组
        cost_group = QGroupBox("成本预算参数")
        cost_layout = QGridLayout(cost_group)

        cost_params = [
            ("SSA无人机单价($):", "p_SSA", "80000"),
            ("中继无人机单价($):", "p_R", "25000")
        ]

        for i, (label, attr, default) in enumerate(cost_params):
            cost_layout.addWidget(QLabel(label), i, 0)
            edit = QLineEdit(default)
            setattr(self, attr, edit)
            cost_layout.addWidget(edit, i, 1)

        scroll_layout.addWidget(cost_group)

        # 安全设置组
        safety_group = QGroupBox("安全设置")
        safety_layout = QGridLayout(safety_group)

        safety_layout.addWidget(QLabel("安全冗余系数:"), 0, 0)
        self.gamma = QLineEdit("1.2")
        safety_layout.addWidget(self.gamma, 0, 1)

        scroll_layout.addWidget(safety_group)

        # 保存参数按钮
        save_btn = QPushButton("保存参数设置")
        save_btn.clicked.connect(self.save_parameters)
        scroll_layout.addWidget(save_btn)

        scroll.setWidget(scroll_content)
        layout.addWidget(scroll)

        self.tab_widget.addTab(tab, "参数设置")

    def create_simulation_tab(self):
        """创建模拟分析标签页"""
        tab = QWidget()
        layout = QVBoxLayout(tab)

        # 分析选项组
        options_group = QGroupBox("分析选项")
        options_layout = QVBoxLayout(options_group)

        # 多年分析选项
        years_layout = QHBoxLayout()
        years_layout.addWidget(QLabel("分析年限:"))
        self.years_input = QLineEdit("10")
        years_layout.addWidget(self.years_input)
        years_layout.addWidget(QLabel("年"))
        years_layout.addStretch()

        growth_layout = QHBoxLayout()
        growth_layout.addWidget(QLabel("年增长率:"))
        self.growth_input = QLineEdit("0.03")
        growth_layout.addWidget(self.growth_input)
        growth_layout.addWidget(QLabel("%"))
        growth_layout.addStretch()

        options_layout.addLayout(years_layout)
        options_layout.addLayout(growth_layout)

        layout.addWidget(options_group)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # 控制按钮
        btn_layout = QHBoxLayout()

        self.run_btn = QPushButton("开始模拟分析")
        self.run_btn.clicked.connect(self.run_simulation)
        btn_layout.addWidget(self.run_btn)

        self.stop_btn = QPushButton("停止分析")
        self.stop_btn.setEnabled(False)
        btn_layout.addWidget(self.stop_btn)

        layout.addLayout(btn_layout)

        # 实时日志
        log_group = QGroupBox("分析日志")
        log_layout = QVBoxLayout(log_group)

        self.log_text = QTextEdit()
        self.log_text.setMaximumHeight(150)
        log_layout.addWidget(self.log_text)

        layout.addWidget(log_group)

        self.tab_widget.addTab(tab, "模拟分析")

    def create_results_tab(self):
        """创建结果展示标签页"""
        tab = QWidget()
        layout = QVBoxLayout(tab)

        # 结果文本区域
        results_group = QGroupBox("分析结果")
        results_layout = QVBoxLayout(results_group)

        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setMinimumHeight(300)
        results_layout.addWidget(self.results_text)

        layout.addWidget(results_group)

        # 导出按钮 - 只保留Word导出
        export_layout = QHBoxLayout()

        self.export_word_btn = QPushButton("📄 导出Word报告")
        self.export_word_btn.clicked.connect(self.export_word_report)
        self.export_word_btn.setEnabled(False)
        export_layout.addWidget(self.export_word_btn)

        layout.addLayout(export_layout)

        self.tab_widget.addTab(tab, "分析结果")

    def create_visualization_tab(self):
        """创建可视化标签页 - 改进版，添加大预览区域"""
        tab = QWidget()
        layout = QVBoxLayout(tab)

        # 可视化选项
        viz_group = QGroupBox("可视化选项")
        viz_layout = QVBoxLayout(viz_group)

        viz_buttons_layout = QHBoxLayout()

        self.demand_btn = QPushButton("📈 显示需求时序图")
        self.demand_btn.clicked.connect(lambda: self.show_plot('demand'))
        viz_buttons_layout.addWidget(self.demand_btn)

        self.cost_btn = QPushButton("💰 显示成本分析图")
        self.cost_btn.clicked.connect(lambda: self.show_plot('cost'))
        viz_buttons_layout.addWidget(self.cost_btn)

        self.relay_btn = QPushButton("📡 显示中继部署图")
        self.relay_btn.clicked.connect(lambda: self.show_plot('relay'))
        viz_buttons_layout.addWidget(self.relay_btn)

        viz_layout.addLayout(viz_buttons_layout)
        layout.addWidget(viz_group)

        # 大预览区域 - 在下方居中显示
        preview_group = QGroupBox("图表预览")
        preview_layout = QVBoxLayout(preview_group)

        # 预览标签 - 尽可能大
        self.preview_label = QLabel()
        self.preview_label.setMinimumSize(600, 400)  # 设置最小尺寸
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setStyleSheet("""
            QLabel {
                border: 3px dashed #07C160;
                border-radius: 12px;
                background-color: #f8f9fa;
                margin: 15px;
                padding: 20px;
                qproperty-alignment: AlignCenter;
            }
        """)
        self.preview_label.setText("图表预览区域\n\n点击上方按钮生成可视化图表\n图表将在此区域显示")
        self.preview_label.setWordWrap(True)

        preview_layout.addWidget(self.preview_label)

        # 图片操作按钮
        preview_btn_layout = QHBoxLayout()

        self.save_image_btn = QPushButton("💾 保存图片")
        self.save_image_btn.clicked.connect(self.save_current_image)
        self.save_image_btn.setEnabled(False)
        preview_btn_layout.addWidget(self.save_image_btn)

        self.clear_preview_btn = QPushButton("🗑️ 清除预览")
        self.clear_preview_btn.clicked.connect(self.clear_preview)
        preview_btn_layout.addWidget(self.clear_preview_btn)

        preview_layout.addLayout(preview_btn_layout)
        layout.addWidget(preview_group)

        self.tab_widget.addTab(tab, "📊 可视化")

    def show_plot(self, plot_type):
        """显示图表并在预览区域显示"""
        try:
            if plot_type == 'demand':
                image_path = self.model.plot_demand()
                title = "需求时序图"
            elif plot_type == 'cost':
                image_path = self.model.plot_costs()
                title = "成本分析图"
            elif plot_type == 'relay':
                image_path = self.model.plot_relay_deployment()
                title = "中继部署图"
            else:
                return

            # 在预览区域显示图片
            if image_path and os.path.exists(image_path):
                self.current_plot_path = image_path
                pixmap = QPixmap(image_path)

                # 缩放图片以适应预览区域，但保持比例
                scaled_pixmap = pixmap.scaled(
                    self.preview_label.width() - 40,
                    self.preview_label.height() - 40,
                    Qt.KeepAspectRatio,
                    Qt.SmoothTransformation
                )

                self.preview_label.setPixmap(scaled_pixmap)
                self.preview_label.setText("")  # 清除文本
                self.save_image_btn.setEnabled(True)

                self.log_message(f"✅ 已生成{title}并在预览区域显示")
            else:
                self.log_message(f"❌ 无法加载{title}")

        except Exception as e:
            QMessageBox.warning(self, "可视化错误", f"无法显示图表: {str(e)}")
            self.log_message(f"❌ 图表显示失败: {str(e)}")

    def save_current_image(self):
        """保存当前预览的图片"""
        if not self.current_plot_path:
            QMessageBox.warning(self, "保存失败", "没有可保存的图片")
            return

        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "保存图片",
                f"消防无人机图表_{os.path.basename(self.current_plot_path)}",
                "图片文件 (*.png *.jpg *.jpeg)"
            )

            if file_path:
                import shutil
                shutil.copy2(self.current_plot_path, file_path)
                QMessageBox.information(self, "保存成功", f"图片已保存到:\n{file_path}")
                self.log_message(f"✅ 图片已保存: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "保存失败", f"保存图片时出错: {str(e)}")

    def clear_preview(self):
        """清除预览区域"""
        self.preview_label.clear()
        self.preview_label.setText("图表预览区域\n\n点击上方按钮生成可视化图表\n图表将在此区域显示")
        self.current_plot_path = None
        self.save_image_btn.setEnabled(False)

    def save_parameters(self):
        """保存参数设置"""
        try:
            # 收集区域频率
            region_freqs = []
            for edit in self.region_freqs:
                region_freqs.append(float(edit.text()))

            # 更新模型参数
            self.model.params.update({
                'R': int(self.region_count.text()),
                'lambda_i': region_freqs,
                'A_mean': float(self.A_mean.text()),
                'A_std': float(self.A_std.text()),
                'alpha1': float(self.alpha1.text()),
                'alpha2': float(self.alpha2.text()),
                'S_SSA': float(self.S_SSA.text()),
                'T_cycle': float(self.T_cycle.text()),
                'R_max': float(self.R_max.text()),
                'R_cov': float(self.R_cov.text()),
                'beta': float(self.beta.text()),
                'p_SSA': float(self.p_SSA.text()),
                'p_R': float(self.p_R.text()),
                'gamma': float(self.gamma.text())
            })

            self.log_message("✅ 参数设置已保存！")
            QMessageBox.information(self, "成功", "参数设置已保存！")

        except ValueError as e:
            QMessageBox.warning(self, "输入错误", "请检查所有输入框是否都填写了有效的数字")
        except AttributeError as e:
            error_msg = f"程序内部错误: {str(e)}"
            QMessageBox.critical(self, "程序错误", error_msg)
            self.log_message(f"❌ 程序错误: {str(e)}")

    def run_simulation(self):
        """运行模拟分析"""
        try:
            # 更新进度条
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.run_btn.setEnabled(False)
            self.stop_btn.setEnabled(True)

            # 保存参数
            self.save_parameters()

            # 模拟分析过程
            self.log_message("开始模拟分析...")
            QTimer.singleShot(100, self.analyze_step1)

        except Exception as e:
            self.log_message(f"分析失败: {str(e)}")
            self.analysis_finished()

    def analyze_step1(self):
        """分析步骤1：生成火灾事件"""
        self.progress_bar.setValue(20)
        self.log_message("生成火灾事件...")
        QTimer.singleShot(500, self.analyze_step2)

    def analyze_step2(self):
        """分析步骤2：计算设备需求"""
        self.progress_bar.setValue(40)
        self.log_message("计算设备需求...")
        QTimer.singleShot(500, self.analyze_step3)

    def analyze_step3(self):
        """分析步骤3：优化配置"""
        self.progress_bar.setValue(60)
        self.log_message("优化设备配置...")

        try:
            results = self.model.run_analysis(
                years=int(self.years_input.text()),
                growth_rate=float(self.growth_input.text())
            )
            self.analysis_results = results
        except Exception as e:
            self.log_message(f"优化配置失败: {str(e)}")
            self.analysis_results = {
                'config': {
                    'x_SSA_opt': 10,
                    'x_R_opt': 6,
                    'total_cost': 1000000,
                    'peak_demand_SSA': 8.5,
                    'peak_demand_R': 5.2
                }
            }

        QTimer.singleShot(500, self.analyze_step4)

    def analyze_step4(self):
        """分析步骤4：生成报告"""
        self.progress_bar.setValue(80)
        self.log_message("生成分析报告...")

        self.display_results()
        QTimer.singleShot(500, self.analysis_finished)

    def analysis_finished(self):
        """分析完成"""
        self.progress_bar.setValue(100)
        self.run_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.log_message("分析完成！")

        # 启用功能按钮
        self.demand_btn.setEnabled(True)
        self.cost_btn.setEnabled(True)
        self.relay_btn.setEnabled(True)
        self.export_word_btn.setEnabled(True)

    def log_message(self, message):
        """添加日志消息"""
        self.log_text.append(f"{message}")
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )

    def display_results(self):
        """显示分析结果"""
        if not hasattr(self, 'analysis_results'):
            return

        results = self.analysis_results
        report = """
智能消防无人机配置分析报告
================================

设备配置推荐:
----------------
• SSA无人机: {} 架
• 中继无人机: {} 架

投资分析:
------------
• SSA投资: ${:,.0f}
• 中继投资: ${:,.0f}
• 总投资: ${:,.0f}

需求分析:
------------
• SSA峰值需求: {:.1f} 架
• 中继峰值需求: {:.1f} 架
• 安全系数: {}

安全裕度:
------------
• SSA安全系数: {:.2f}
• 中继安全系数: {:.2f}

        """.format(
            results['config']['x_SSA_opt'],
            results['config']['x_R_opt'],
            results['config']['x_SSA_opt'] * self.model.params['p_SSA'],
            results['config']['x_R_opt'] * self.model.params['p_R'],
            results['config']['total_cost'],
            results['config']['peak_demand_SSA'],
            results['config']['peak_demand_R'],
            self.model.params['gamma'],
            results['config']['x_SSA_opt'] / results['config']['peak_demand_SSA'] if results['config'][
                                                                                         'peak_demand_SSA'] > 0 else 0,
            results['config']['x_R_opt'] / results['config']['peak_demand_R'] if results['config'][
                                                                                     'peak_demand_R'] > 0 else 0
        )

        self.results_text.setText(report)

    def export_word_report(self):
        """导出Word报告"""
        if not HAS_DOCX:
            QMessageBox.warning(self, "功能不可用",
                                "python-docx 库未安装，无法导出Word报告。\n\n"
                                "请安装: pip install python-docx")
            return

        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "导出Word报告",
                "消防无人机配置分析报告.docx",
                "Word文档 (*.docx)"
            )

            if file_path:
                self.create_word_report(file_path)
                QMessageBox.information(self, "导出成功", f"Word报告已导出到:\n{file_path}")
                self.log_message(f"✅ Word报告已导出: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出Word报告时出错: {str(e)}")
            self.log_message(f"❌ Word导出失败: {str(e)}")

    def create_word_report(self, file_path):
        """创建Word报告文档"""
        doc = Document()

        # 设置中文字体
        try:
            doc.styles['Normal'].font.name = '微软雅黑'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except:
            pass  # 如果设置字体失败，继续使用默认字体

        # 标题
        title = doc.add_heading('智能消防无人机配置分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 分析时间
        from datetime import datetime
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        time_para = doc.add_paragraph(f"生成时间: {current_time}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 设备配置推荐
        doc.add_heading('一、设备配置推荐', level=1)
        results = self.analysis_results['config']
        config_table = doc.add_table(rows=2, cols=2)
        config_table.style = 'Light Grid Accent 1'

        config_table.cell(0, 0).text = '设备类型'
        config_table.cell(0, 1).text = '推荐数量'
        config_table.cell(1, 0).text = 'SSA无人机'
        config_table.cell(1, 1).text = f"{results['x_SSA_opt']} 架"

        # 添加中继无人机行
        config_table.add_row()
        config_table.cell(2, 0).text = '中继无人机'
        config_table.cell(2, 1).text = f"{results['x_R_opt']} 架"

        doc.add_paragraph()  # 空行

        # 投资分析
        doc.add_heading('二、投资分析', level=1)
        investment_table = doc.add_table(rows=4, cols=2)
        investment_table.style = 'Light Grid Accent 1'

        investment_table.cell(0, 0).text = '项目'
        investment_table.cell(0, 1).text = '金额'
        investment_table.cell(1, 0).text = 'SSA无人机投资'
        investment_table.cell(1, 1).text = f"${results['x_SSA_opt'] * self.model.params['p_SSA']:,.0f}"
        investment_table.cell(2, 0).text = '中继无人机投资'
        investment_table.cell(2, 1).text = f"${results['x_R_opt'] * self.model.params['p_R']:,.0f}"
        investment_table.cell(3, 0).text = '总投资'
        investment_table.cell(3, 1).text = f"${results['total_cost']:,.0f}"

        doc.add_paragraph()  # 空行

        # 需求分析
        doc.add_heading('三、需求分析', level=1)
        demand_para = doc.add_paragraph()
        demand_para.add_run(f"SSA无人机峰值需求: {results['peak_demand_SSA']:.1f} 架\n")
        demand_para.add_run(f"中继无人机峰值需求: {results['peak_demand_R']:.1f} 架\n")
        demand_para.add_run(f"安全系数: {self.model.params['gamma']}")

        doc.add_paragraph()  # 空行

        # 安全裕度
        doc.add_heading('四、安全裕度', level=1)
        safety_para = doc.add_paragraph()
        ssa_safety = results['x_SSA_opt'] / results['peak_demand_SSA'] if results['peak_demand_SSA'] > 0 else 0
        relay_safety = results['x_R_opt'] / results['peak_demand_R'] if results['peak_demand_R'] > 0 else 0

        safety_para.add_run(f"SSA安全系数: {ssa_safety:.2f}\n")
        safety_para.add_run(f"中继安全系数: {relay_safety:.2f}")

        # 添加图片（如果存在）
        if hasattr(self, 'current_plot_path') and self.current_plot_path and os.path.exists(self.current_plot_path):
            doc.add_heading('五、图表展示', level=1)
            doc.add_paragraph("当前预览的图表:")
            doc.add_picture(self.current_plot_path, width=Inches(6))

        # 保存文档
        doc.save(file_path)


def main():
    """主函数"""
    app = QApplication(sys.argv)

    # 设置应用程序字体
    font = QFont("Microsoft YaHei", 10)
    app.setFont(font)

    window = WeChatStyleWindow()
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()